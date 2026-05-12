import streamlit as st
import random
import os
import re
import datetime
from copy import deepcopy
from io import BytesIO
from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Pt

# --- AYARLAR VE HAFIZA SİSTEMİ ---
SABLON_YOLU = "manuel.docx"  # Klasördeki dosya adı
CIHAZ_DOSYASI = "cihazlar.txt"

# BULUT İÇİN: Masaüstü yolunu sildik, Streamlit hafızasını kullanacağız.

def cihazlari_yukle():
    varsayilan = ["DFA – 0469 KUMPAS", "DFA – 3262 MİHENGİR", "M3 DİŞ MASTARI", "M5 DİŞ MASTARI"]
    if os.path.exists(CIHAZ_DOSYASI):
        with open(CIHAZ_DOSYASI, "r", encoding="utf-8") as f:
            return sorted(list(set(varsayilan + [line.strip() for line in f.readlines() if line.strip()])))
    return sorted(varsayilan)

def cihaz_kaydet(yeni_cihaz):
    if yeni_cihaz and yeni_cihaz != "Yeni Cihaz Ekle...":
        mevcut = cihazlari_yukle()
        if yeni_cihaz not in mevcut:
            try:
                with open(CIHAZ_DOSYASI, "a", encoding="utf-8") as f:
                    f.write(yeni_cihaz.strip() + "\n")
            except: pass # Bulutta dosya yazma bazen kısıtlıdır, hata vermesin

def hucre_yaz(hucre, metin):
    hucre.text = str(metin)
    if hucre.paragraphs:
        p = hucre.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        if p.runs:
            p.runs[0].font.name = 'Calibri Light'
            p.runs[0].font.size = Pt(9)

def benzersiz_hucreler(satir):
    res = []
    for c in satir.cells:
        if c not in res: res.append(c)
    return res

def parametre_tetikleyici(i):
    n_key, c_key, ta_key, te_key = f"n_{i}", f"c_box_{i}", f"ta_{i}", f"te_{i}"
    if n_key in st.session_state:
        v = st.session_state[n_key].upper().strip()
        if any(x in v for x in ["DİŞ", "DIS", "HELI", "METRİK", "METRIK"]) or (v.startswith('M') and any(c.isdigit() for c in v)):
            match = re.search(r'\d+', v)
            m_num = match.group() if match else ""
            if m_num:
                tahmin = f"M{m_num} HELICOIL DİŞ MASTARI" if "HELI" in v else f"M{m_num} DİŞ MASTARI"
                st.session_state[c_key] = tahmin
                st.session_state[ta_key] = "GO"
                st.session_state[te_key] = "NO GO"

# --- ARAYÜZ ---
st.set_page_config(page_title="Teknolus Manuel Ölçü Paneli", layout="wide")
st.sidebar.header("📋 Belge Üst Bilgileri")

numune_sayisi = st.sidebar.number_input("Numune (Satır) Sayısı:", min_value=1, value=10)
siparis = st.sidebar.text_input("Sipariş No (Envanter):")
tarih = st.sidebar.text_input("Tarih:", datetime.date.today().strftime("%d.%m.%Y"))
m_no = st.sidebar.text_input("Malzeme No:")
m_adi = st.sidebar.text_input("Malzeme Açıklaması:")
g_mik = st.sidebar.text_input("Gelen Miktar:")
k_mik = st.sidebar.text_input("Kontrol Miktarı:", value=str(numune_sayisi))
u_mik = st.sidebar.text_input("Uygun Miktar:", value=str(numune_sayisi))
k_eden = st.sidebar.text_input("Kontrol Eden:")
onay = st.sidebar.text_input("Onaylayan:")

st.subheader("📏 Manuel Ölçüm Parametreleri")
ayarlar = []
mevcut_list = cihazlari_yukle() + ["Yeni Cihaz Ekle..."]

for i in range(1, 10):
    with st.expander(f"Ölçüm Noktası {i}", expanded=(i<=3)):
        c1, c2, c3, c4, c5, c6 = st.columns([0.8, 1.2, 1, 1, 2.5, 1.5])
        b_no = c1.text_input("Balon No", key=f"b_{i}") 
        nom = c2.text_input("Nominal", key=f"n_{i}", on_change=parametre_tetikleyici, args=(i,))
        ta = c3.text_input("Tol(+)", key=f"ta_{i}")
        te = c4.text_input("Tol(-)", key=f"te_{i}")
        c_sec = c5.selectbox("Cihaz Seç", mevcut_list, key=f"c_box_{i}")
        f_cihaz = c_sec
        if c_sec == "Yeni Cihaz Ekle...":
            f_cihaz = c6.text_input("Cihaz Adı:", key=f"c_man_{i}")
        ayarlar.append({'b': b_no, 'n': nom, 'ta': ta, 'te': te, 'c': f_cihaz, 'aktif': bool(b_no.strip())})

if st.button("🚀 FORMU OLUŞTUR", type="primary"):
    try:
        tpl = DocxTemplate(SABLON_YOLU)
        ctx = {
            'envanter_no': siparis, 'tarih': tarih, 'malzeme_no': m_no, 'malzeme_adi': m_adi,
            'genel_miktar': g_mik, 'kontrol_miktari': k_mik, 'uygun_miktar': u_mik,
            'kontrol_eden': k_eden, 'onaylayan': onay
        }
        
        for i, a in enumerate(ayarlar, 1):
            if a['aktif']:
                ctx.update({f'balonno{i}': a['b'], f'nom{i}': a['n'], f'tolarti{i}': a['ta'], f'toleksi{i}': a['te']})
                ctx[f'cihaz{i if i > 1 else 2}'] = a['c']
            else:
                ctx.update({f'balonno{i}': '', f'nom{i}': '', f'tolarti{i}': '', f'toleksi{i}': '', f'cihaz{i if i > 1 else 2}': ''})

        tpl.render(ctx)
        out_tmp = BytesIO()
        tpl.save(out_tmp)
        doc = Document(BytesIO(out_tmp.getvalue()))
        tablo = doc.tables[0]
        ornek_satir_xml = tablo.rows[8]._tr

        for s in range(numune_sayisi):
            if s < 10: row = tablo.rows[8 + s]
            else:
                yeni_tr = deepcopy(ornek_satir_xml)
                tablo.rows[8 + s]._tr.addprevious(yeni_tr)
                row = tablo.rows[8 + s]
            h = benzersiz_hucreler(row)
            hucre_yaz(h[0], str(s+1)); hucre_yaz(h[1], f"{s+1:02d}")
            for i, a in enumerate(ayarlar):
                v_idx, r_idx = 2+(i*2), 3+(i*2)
                if a['aktif'] and v_idx < len(h):
                    if "GO" in str(a['ta']).upper() or any(x in str(a['c']).upper() for x in ["DİŞ", "HELI"]):
                        hucre_yaz(h[v_idx], "GO")
                    else:
                        try:
                            n_str = str(a['n']).replace(',', '.')
                            ta_str = str(a['ta']).replace(',', '.')
                            te_str = str(a['te']).replace(',', '.')
                            
                            # Girilen tolerans değerlerindeki virgülden sonraki basamak sayısını hesaplıyoruz
                            def basamak_sayisi(metin):
                                return len(metin.split('.')[1]) if '.' in metin else 0
                                
                            # + ve - toleranslardan hangisinin hassasiyeti (basamak sayısı) daha yüksekse onu baz al
                            hassasiyet = max(basamak_sayisi(ta_str), basamak_sayisi(te_str))
                            
                            # Eğer toleranslar tam sayı girilmişse varsayılan olarak kumpas hassasiyeti (2 basamak) kalsın
                            if hassasiyet == 0: 
                                hassasiyet = 2 

                            n = float(n_str)
                            ust = n + float(ta_str)
                            e_s = te_str
                            alt = n + float(e_s) if e_s.startswith('+') else n - float(e_s or 0)
                            
                            # Rastgele değeri üret ve dinamik hassasiyete göre formatla (örn: hassasiyet 3 ise 5.120 yazar)
                            uretilen_deger = random.uniform(alt, ust)
                            formatli_deger = f"{uretilen_deger:.{hassasiyet}f}"
                            
                            hucre_yaz(h[v_idx], formatli_deger)
                        except: hucre_yaz(h[v_idx], "GO")
                    hucre_yaz(h[r_idx], "OK")
                elif v_idx < len(h):
                    hucre_yaz(h[v_idx], ""); hucre_yaz(h[r_idx], "")

        if numune_sayisi < 10:
            for s in range(numune_sayisi, 10):
                for cell in benzersiz_hucreler(tablo.rows[8+s]): hucre_yaz(cell, "")

        # --- İNDİRME SİSTEMİ ---
        final_out = BytesIO()
        doc.save(final_out)
        final_out.seek(0)
        
        st.success("✅ Form başarıyla oluşturuldu!")
        st.download_button(
            label="📥 Manuel Ölçü Formunu İndir",
            data=final_out,
            file_name=f"MANUEL_{m_no}_{tarih}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    except Exception as e:
        st.error(f"Hata: {e}")
